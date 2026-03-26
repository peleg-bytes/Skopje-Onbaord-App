"""One-off generator: Skopje Onboard Survey user guide (.docx). Run: py -3 scripts/build_user_guide_docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def screenshot_box(doc: Document, title: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Screenshot placeholder")
    r.bold = True
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(title)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p3 = cell.add_paragraph("Paste or insert your image here (replace this text).")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    t.rows[0].height = Inches(2.2)
    doc.add_paragraph()


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "docs" / "Skopje-Onboard-User-Guide.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph("Skopje Onboard Survey")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(22)
    doc.add_paragraph("Short user guide — how to use the app and view submissions.")
    doc.add_paragraph()

    doc.add_heading("1. First launch & home screen", level=1)
    doc.add_paragraph(
        "Open the app. On the main screen you choose the station and surveyor name, then use the buttons to sync, start a survey, open submissions, or settings."
    )
    screenshot_box(doc, "Home screen — station, surveyor, Sync now, Start survey")
    screenshot_box(doc, "Top bar — Submissions and Settings")

    doc.add_heading("2. Syncing with the server", level=1)
    doc.add_paragraph(
        "Tap Sync now to download the latest survey setup from the server. "
        "Wait until it finishes before starting a survey if you need the newest configuration."
    )
    screenshot_box(doc, "Sync in progress or success message (toast or status)")

    doc.add_heading("3. During a survey (counting)", level=1)
    doc.add_paragraph(
        "Tap Start survey. Use + and − to record counts. GPS and server status may appear on screen. "
        "Submit survey sends your answers; Reset clears the current counts (use carefully)."
    )
    screenshot_box(doc, "Counting screen with totals and Submit / Reset")

    doc.add_heading("4. Submitting", level=1)
    doc.add_paragraph(
        "After Submit survey, the app sends data when online. If there is no connection, the submission may be queued for later. "
        "Success or error messages are shown as short on-screen notices (toasts)."
    )
    screenshot_box(doc, "Success or queued submission message")

    doc.add_heading("5. Viewing submissions", level=1)
    doc.add_paragraph(
        "Tap Submissions (top bar) to open the list of local submissions. "
        "Each row shows status and details. Use Try sync to server now (or equivalent) to retry sending queued items when you have network."
    )
    screenshot_box(doc, "Submissions list with status banner")
    screenshot_box(doc, "Submission detail or retry sync action")

    doc.add_heading("6. Settings & language", level=1)
    doc.add_paragraph(
        "Open Settings to adjust options such as language (e.g. Macedonian / English). "
        "Some messages follow the language you select."
    )
    screenshot_box(doc, "Settings screen — language and other options")

    doc.add_heading("7. Tips", level=1)
    doc.add_paragraph(
        "• Allow location permission if the survey requires GPS.\n"
        "• Use Wi‑Fi or mobile data for sync and submit.\n"
        "• If a submit fails, check Submissions and try sync again when online."
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
